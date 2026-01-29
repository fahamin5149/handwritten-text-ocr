"""
OCR Engine using Microsoft TrOCR and Google Gemini to extract text from handwritten and printed images.
Converts images to structured Word documents with preserved formatting.
"""

import os
import sys
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import cv2
import numpy as np
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import torch
from google import genai
from google.genai import types
from dotenv import load_dotenv

from utils.image_processing import preprocess_image

# Load environment variables
load_dotenv()


# TrOCR Model Configuration
# Using handwritten model for better handwriting recognition
TROCR_MODEL_NAME = "microsoft/trocr-base-handwritten"
# Alternative models:
# "microsoft/trocr-large-handwritten" - Better quality but slower
# "microsoft/trocr-base-printed" - For printed text only


@dataclass
class TextRegion:
    """Represents a text region extracted from the image."""
    text: str
    x: int
    y: int
    width: int
    height: int
    is_paragraph_start: bool = False


class OCRConverter:
    """
    Main OCR conversion class that handles image-to-Word conversion using TrOCR or Gemini.
    """
    
    def __init__(self, ocr_engine: str = "microsoft"):
        """Initialize OCR converter with specified engine.
        
        Args:
            ocr_engine: Either "microsoft" for TrOCR or "gemini" for Google Gemini
        """
        self.ocr_engine = ocr_engine.lower()
        self.processor = None
        self.model = None
        self.gemini_client = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        if self.ocr_engine == "microsoft":
            self._load_model()
        elif self.ocr_engine == "gemini":
            self._load_gemini_client()
        else:
            raise ValueError(f"Unknown OCR engine: {ocr_engine}. Use 'microsoft' or 'gemini'.")
    
    def _load_model(self) -> None:
        """Load TrOCR model and processor."""
        try:
            print(f"Loading TrOCR model ({TROCR_MODEL_NAME}) on {self.device}...")
            self.processor = TrOCRProcessor.from_pretrained(TROCR_MODEL_NAME)
            self.model = VisionEncoderDecoderModel.from_pretrained(TROCR_MODEL_NAME)
            self.model.to(self.device)
            self.model.eval()
            print("Model loaded successfully!")
        except Exception as e:
            raise RuntimeError(
                f"Failed to load TrOCR model.\n\n"
                f"Error: {str(e)}\n\n"
                "Please ensure you have installed all dependencies:\n"
                "pip install transformers torch pillow\n\n"
                "The model will be downloaded automatically on first run."
            )
    
    def _load_gemini_client(self) -> None:
        """Load Gemini API client."""
        try:
            # Try to get API key from Streamlit secrets first (for cloud deployment)
            api_key = None
            try:
                import streamlit as st
                if hasattr(st, 'secrets') and 'GEMINI_API_KEY' in st.secrets:
                    api_key = st.secrets['GEMINI_API_KEY']
                    print("Using Gemini API key from Streamlit secrets")
            except (ImportError, FileNotFoundError):
                pass
            
            # Fall back to environment variable (for local development)
            if not api_key:
                api_key = os.getenv("GEMINI_API_KEY")
                if api_key:
                    print("Using Gemini API key from .env file")
            
            if not api_key or api_key == "your_gemini_api_key_here":
                raise ValueError(
                    "Gemini API key not found or not configured.\n\n"
                    "For local development: Set GEMINI_API_KEY in your .env file\n"
                    "For Streamlit Cloud: Add GEMINI_API_KEY to Streamlit secrets\n"
                    "Get your API key from: https://aistudio.google.com/app/apikey"
                )
            
            print("Initializing Gemini client...")
            self.gemini_client = genai.Client(api_key=api_key)
            print("Gemini client initialized successfully!")
        except Exception as e:
            raise RuntimeError(
                f"Failed to initialize Gemini client.\n\n"
                f"Error: {str(e)}\n\n"
                "Please ensure you have:\n"
                "1. Set GEMINI_API_KEY in Streamlit secrets (cloud) or .env file (local)\n"
                "2. Installed google-genai: pip install google-genai"
            )
    
    def convert_image_to_docx(
        self, 
        image_path: str, 
        output_path: str,
        preprocess: bool = False,
        progress_callback=None
    ) -> bool:
        """
        Convert an image to a formatted Word document using TrOCR.
        
        Args:
            image_path: Path to input image file
            output_path: Path where .docx file will be saved
            preprocess: Whether to apply image preprocessing (less critical for TrOCR)
            progress_callback: Optional callback function to report progress
            
        Returns:
            True if conversion successful, False otherwise
        """
        try:
            # Load image
            if progress_callback:
                progress_callback("Loading image...")
            image = Image.open(image_path)
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Detect text regions in the image
            if progress_callback:
                progress_callback("Detecting text regions...")
            text_regions = self._detect_text_regions(image_path, preprocess)
            
            # Extract text from each region using selected OCR engine
            if self.ocr_engine == "gemini":
                # For Gemini, process the entire image at once
                if progress_callback:
                    progress_callback("Extracting text with Gemini...")
                full_text = self._extract_text_with_gemini(image)
                # Create a single region with the full text
                if text_regions:
                    text_regions[0].text = full_text
                    text_regions = [text_regions[0]]  # Use only the first region
            else:
                # For TrOCR, process each region separately
                total_regions = len(text_regions)
                for idx, region in enumerate(text_regions):
                    if progress_callback:
                        progress_callback(f"Extracting text from region {idx + 1}/{total_regions}...")
                    
                    region_image = image.crop((
                        region.x, 
                        region.y, 
                        region.x + region.width, 
                        region.y + region.height
                    ))
                    region.text = self._extract_text_with_trocr(region_image)
            
            # Generate Word document
            if progress_callback:
                progress_callback("Creating Word document...")
            self._create_word_document(text_regions, output_path)
            
            if progress_callback:
                progress_callback("Complete!")
            
            return True
            
        except Exception as e:
            print(f"Error during conversion: {str(e)}")
            raise
    
    def _detect_text_regions(self, image_path: str, preprocess: bool) -> List[TextRegion]:
        """
        Detect text regions in the image using OpenCV.
        
        Args:
            image_path: Path to input image
            preprocess: Whether to preprocess the image
            
        Returns:
            List of TextRegion objects
        """
        # Load image with OpenCV
        if preprocess:
            img = preprocess_image(image_path)
        else:
            img = cv2.imread(image_path)
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding
        _, binary = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        
        # Find contours
        contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        # Sort contours by vertical position (top to bottom)
        regions = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            
            # Filter out very small regions (noise)
            if w > 20 and h > 10:
                regions.append(TextRegion(
                    text="",
                    x=x,
                    y=y,
                    width=w,
                    height=h
                ))
        
        # Sort regions by Y position (top to bottom), then X (left to right)
        regions.sort(key=lambda r: (r.y, r.x))
        
        # Merge overlapping or nearby regions
        merged_regions = self._merge_regions(regions)
        
        # Detect paragraph breaks based on vertical spacing
        if len(merged_regions) > 1:
            for i in range(1, len(merged_regions)):
                prev_region = merged_regions[i-1]
                curr_region = merged_regions[i]
                
                # If vertical gap is significant, mark as paragraph start
                vertical_gap = curr_region.y - (prev_region.y + prev_region.height)
                if vertical_gap > 30:  # Threshold for paragraph break
                    curr_region.is_paragraph_start = True
        
        if merged_regions:
            merged_regions[0].is_paragraph_start = True
        
        return merged_regions
    
    def _merge_regions(self, regions: List[TextRegion]) -> List[TextRegion]:
        """
        Merge nearby text regions that likely belong to the same line.
        
        Args:
            regions: List of text regions
            
        Returns:
            Merged list of text regions
        """
        if not regions:
            return []
        
        merged = []
        current_line = [regions[0]]
        
        for i in range(1, len(regions)):
            prev = current_line[-1]
            curr = regions[i]
            
            # Check if regions are on the same line (similar Y position)
            y_overlap = abs(prev.y - curr.y) < min(prev.height, curr.height) * 0.5
            
            # Check if regions are horizontally close
            x_distance = curr.x - (prev.x + prev.width)
            x_close = x_distance < prev.width * 2
            
            if y_overlap and x_close:
                current_line.append(curr)
            else:
                # Merge current line into one region
                if current_line:
                    merged.append(self._merge_line_regions(current_line))
                current_line = [curr]
        
        # Don't forget the last line
        if current_line:
            merged.append(self._merge_line_regions(current_line))
        
        return merged
    
    def _merge_line_regions(self, line_regions: List[TextRegion]) -> TextRegion:
        """Merge multiple regions from the same line into one."""
        if len(line_regions) == 1:
            return line_regions[0]
        
        x_min = min(r.x for r in line_regions)
        y_min = min(r.y for r in line_regions)
        x_max = max(r.x + r.width for r in line_regions)
        y_max = max(r.y + r.height for r in line_regions)
        
        return TextRegion(
            text="",
            x=x_min,
            y=y_min,
            width=x_max - x_min,
            height=y_max - y_min
        )
    
    def _extract_text_with_trocr(self, image: Image.Image) -> str:
        """
        Extract text from an image using TrOCR.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text string
        """
        try:
            # Preprocess image for TrOCR
            pixel_values = self.processor(image, return_tensors="pt").pixel_values
            pixel_values = pixel_values.to(self.device)
            
            # Generate text
            with torch.no_grad():
                generated_ids = self.model.generate(pixel_values)
            
            # Decode the generated text
            generated_text = self.processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
            
            return generated_text.strip()
            
        except Exception as e:
            print(f"Error extracting text with TrOCR: {str(e)}")
            return ""
    
    def _extract_text_with_gemini(self, image: Image.Image) -> str:
        """
        Extract text from an image using Google Gemini.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text string
        """
        try:
            # Convert PIL Image to bytes
            import io
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_byte_arr = img_byte_arr.getvalue()
            
            # Create image part
            image_part = types.Part.from_bytes(
                data=img_byte_arr,
                mime_type='image/png'
            )
            
            # Create prompt for OCR
            prompt = (
                "Extract all the text from this image. "
                "Preserve the layout and structure as much as possible. "
                "Return only the extracted text, without any additional commentary or formatting."
            )
            
            # Generate content
            response = self.gemini_client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[prompt, image_part]
            )
            
            return response.text.strip()
            
        except Exception as e:
            print(f"Error extracting text with Gemini: {str(e)}")
            return ""
    
    def _create_word_document(
        self, 
        text_regions: List[TextRegion], 
        output_path: str
    ) -> None:
        """
        Create a Word document from extracted text regions.
        
        Args:
            text_regions: List of TextRegion objects with extracted text
            output_path: Path where .docx file will be saved
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Calculate base indent (minimum left position)
        if text_regions:
            all_x = [r.x for r in text_regions]
            min_x = min(all_x) if all_x else 0
            
            # Process each text region
            current_paragraph_lines = []
            
            for i, region in enumerate(text_regions):
                text = region.text.strip()
                
                if not text:
                    continue
                
                # Start new paragraph if marked or first region
                if region.is_paragraph_start and current_paragraph_lines:
                    # Write accumulated lines as paragraph
                    para_text = '\n'.join(current_paragraph_lines)
                    para = doc.add_paragraph(para_text)
                    para.paragraph_format.space_after = Pt(8)
                    
                    # Set font
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                    
                    current_paragraph_lines = []
                
                # Add text to current paragraph
                current_paragraph_lines.append(text)
            
            # Don't forget the last paragraph
            if current_paragraph_lines:
                para_text = '\n'.join(current_paragraph_lines)
                para = doc.add_paragraph(para_text)
                para.paragraph_format.space_after = Pt(8)
                
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        else:
            # No text found
            para = doc.add_paragraph("No text detected in the image.")
            para.paragraph_format.space_after = Pt(6)
        
        # Save document
        doc.save(output_path)
    
    def get_preview_text(self, image_path: str, max_chars: int = 500) -> str:
        """
        Get a preview of the extracted text without creating a document.
        
        Args:
            image_path: Path to input image file
            max_chars: Maximum number of characters to return
            
        Returns:
            Preview text string
        """
        try:
            image = Image.open(image_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using TrOCR
            text = self._extract_text_with_trocr(image)
            
            if len(text) > max_chars:
                return text[:max_chars] + "..."
            return text
            
        except Exception as e:
            return f"Error extracting text: {str(e)}"
